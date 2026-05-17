"""Render IJNMBE supplementary material as portal-ready DOCX files.

Outputs (under manuscripts/ijnmbe/rendered/supplementary/):
  table_s1_rf_baseline.docx
  table_s2_per_stratum_coverage.docx
  supplementary_information.docx   (cover + table titles + S-file index)

Formatting follows the IJNMBE Wiley Free Format defaults baked into
`_reference.docx`:
  - Times New Roman 12 pt body, 11 pt tables
  - Double spacing, continuous line numbers from line 1
  - A4 page, 2.54 cm margins, upper-right page numbers
  - Table headers bold; alternating row shading off; thin grid borders
  - Table titles ABOVE the table; footnotes BELOW

Sources:
  data/results/supplementary/table_s1_rf_baseline.json
  data/results/supplementary/table_s2_per_stratum_coverage.json
  data/results/cqr/cqr_vs_mondrian_time_to_gloc.json
  data/results/sensitivity/sobol_second_order.csv

Usage:
    cd /root/repos/CAMI-Gz-Effects-Model-CGEM-
    /root/.venvs/cgem-ci/bin/python scripts/render_supplementary_docx.py
"""

from __future__ import annotations

import csv
import json
import math
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

REPO = Path(__file__).resolve().parent.parent
SUP_DATA = REPO / "data" / "results" / "supplementary"
SOBOL2 = REPO / "data" / "results" / "sensitivity" / "sobol_second_order.csv"
# Wiley portal field structure:
#   supporting_info/  → uploaded as "Supporting Information" (DOCX tables, model cards, …)
#   data_files/       → uploaded as "Data Files" (code, raw JSON / CSV / Parquet), deposited
#                       to figshare under CC-Zero on acceptance
OUT_ROOT = REPO / "manuscripts" / "ijnmbe" / "rendered" / "supplementary"
OUT_DIR = OUT_ROOT / "supporting_info"
DATA_DIR = OUT_ROOT / "data_files"
OUT_DIR.mkdir(parents=True, exist_ok=True)
DATA_DIR.mkdir(parents=True, exist_ok=True)


# ──────────────────────────────────────────────────────────────────────
# Page / style helpers
# ──────────────────────────────────────────────────────────────────────


def _apply_page_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)

    # Continuous line numbering
    sectPr = section._sectPr
    lnNumType = OxmlElement("w:lnNumType")
    lnNumType.set(qn("w:countBy"), "1")
    lnNumType.set(qn("w:restart"), "continuous")
    lnNumType.set(qn("w:start"), "1")
    lnNumType.set(qn("w:distance"), "720")
    sectPr.append(lnNumType)

    # Page numbers in header (upper-right)
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def _set_run_font(run, *, bold: bool = False, italic: bool = False, size_pt: float = 12):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    # East-Asian + complex-script fonts
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for tag in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(tag), "Times New Roman")


def _set_paragraph_double_spacing(p, *, before: float = 0, after: float = 0):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def _add_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph()
    size = {1: 14, 2: 13, 3: 12}.get(level, 12)
    run = p.add_run(text)
    _set_run_font(run, bold=True, size_pt=size)
    _set_paragraph_double_spacing(p, before=12, after=6)
    return p


def _add_para(doc, text: str, *, bold: bool = False, italic: bool = False, size_pt: float = 12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, bold=bold, italic=italic, size_pt=size_pt)
    _set_paragraph_double_spacing(p)
    return p


def _add_caption(doc, label: str, body: str):
    """Bold label + roman body, single combined paragraph."""
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    _set_run_font(r1, bold=True, size_pt=11)
    r2 = p.add_run(" " + body)
    _set_run_font(r2, size_pt=11)
    _set_paragraph_double_spacing(p, before=6, after=6)
    return p


def _add_footnote(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run_font(r, size_pt=10)
    _set_paragraph_double_spacing(p, before=3, after=3)
    return p


def _set_cell_borders(cell, thickness: int = 4):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(thickness))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tcBorders.append(b)


def _cell_text(cell, text: str, *, bold: bool = False, size_pt: float = 11,
               align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = para.add_run(text)
    _set_run_font(run, bold=bold, size_pt=size_pt)
    _set_cell_borders(cell)


def _build_table(doc, headers: list[str], rows: list[list[str]], *,
                 col_widths_cm: list[float] | None = None,
                 align_per_col: list[int] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if col_widths_cm:
        for col_idx, w in enumerate(col_widths_cm):
            for cell in table.columns[col_idx].cells:
                cell.width = Cm(w)
    aligns = align_per_col or [WD_ALIGN_PARAGRAPH.CENTER] * len(headers)
    # Header
    for j, h in enumerate(headers):
        _cell_text(table.rows[0].cells[j], h, bold=True, align=aligns[j])
    # Body
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            _cell_text(table.rows[i].cells[j], val, bold=False, align=aligns[j])
    return table


def _fmt_ci(triple) -> str:
    if triple is None:
        return "—"
    point, lo, hi = triple
    if any(map(lambda v: v is None or (isinstance(v, float) and math.isnan(v)), (point, lo, hi))):
        return "—"
    if abs(point) >= 100:
        return f"{point:.1f} [{lo:.1f}, {hi:.1f}]"
    return f"{point:.3f} [{lo:.3f}, {hi:.3f}]"


def _fmt_pct(v) -> str:
    if v is None or (isinstance(v, float) and math.isnan(v)):
        return "—"
    return f"{v:.3f}"


def _fmt_cp(lo, hi) -> str:
    if lo is None or hi is None or (isinstance(lo, float) and math.isnan(lo)):
        return "—"
    return f"[{lo:.3f}, {hi:.3f}]"


# ──────────────────────────────────────────────────────────────────────
# Table S1 — RF baseline
# ──────────────────────────────────────────────────────────────────────


def render_table_s1():
    data = json.loads((SUP_DATA / "table_s1_rf_baseline.json").read_text())
    doc = Document()
    _apply_page_layout(doc)

    _add_heading(doc, "Supplementary Table S1", level=1)
    _add_caption(
        doc,
        "Table S1.",
        "RandomForest baseline vs. XGBoost surrogate on the OSF-pre-registered held-out test "
        "split (seed 42). Continuous targets are evaluated on all test rows (n = 487); "
        "censored time-to-event targets are evaluated on the event-positive subset of the "
        "test split. Point estimates are accompanied by 95 % paired bootstrap CIs "
        f"({data['_meta']['n_boot']} resamples). Bold targets are censored; the "
        "`R² damped` column reports the R² of the expected-time prediction "
        "P(event) × E[time | event] evaluated on event-positive rows — the failure mode "
        "of the naive two-stage strategy that the conditional regressor avoids.",
    )

    headers = [
        "Target",
        "n",
        "RF R² [95 % CI]",
        "RF RMSE [95 % CI]",
        "XGB R² [95 % CI]",
        "XGB RMSE [95 % CI]",
        "RF R² damped",
        "XGB R² damped",
    ]
    rows = []
    aligns = [
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
    ]
    for r in data["rows"]:
        is_cens = r["censored"]
        label = r["target"] + (" (event=1)" if is_cens else "")
        rows.append([
            label,
            str(r["n_test_event"]),
            _fmt_ci(r["rf_r2_event"]),
            _fmt_ci(r["rf_rmse_event"]),
            _fmt_ci(r["xgb_r2_event"]),
            _fmt_ci(r["xgb_rmse_event"]),
            _fmt_ci(r.get("rf_r2_damped_event")) if is_cens else "—",
            _fmt_ci(r.get("xgb_r2_damped_event")) if is_cens else "—",
        ])

    col_widths = [3.6, 1.0, 2.4, 2.4, 2.4, 2.4, 2.0, 2.0]
    _build_table(doc, headers, rows, col_widths_cm=col_widths, align_per_col=aligns)

    _add_footnote(
        doc,
        "Footnote — RF defaults: n_estimators = 400, random_state = 42, no monotonicity "
        "constraints (sklearn RandomForest does not expose them). XGBoost is the production "
        "TwoStageXGBSurrogate with per-target monotonicity priors and the OSF-pre-registered "
        "hyperparameter search space. `R² damped` evaluates the expected-time prediction "
        "P(event) × E[time | event] on event-positive rows. Because P(event) is < 1 on every "
        "row, the prediction is systematically damped below the true event time, producing "
        "large negative R² values; RF's classifier under-predicts P(event=1) on event-positive "
        f"rows more severely than XGBoost's (mean RF P̂ = "
        f"{data['rows'][0]['rf_p_event_mean_on_event_pos']:.3f} vs XGB P̂ = "
        f"{data['rows'][0]['xgb_p_event_mean_on_event_pos']:.3f} for time_to_greyout_s), which "
        "explains the larger negative R² damped for RF. This is the failure mode discussed in "
        "manuscript §3.2 and is the reason the production framework reports stage-2 R² (conditional "
        "on event=1), not the damped product, as the operational accuracy metric for censored "
        "regressors. The RF baseline is reported in this supplementary table for transparency; the "
        "architectural difference (XGB's monotonicity-constrained two-stage path) makes the "
        "comparison apples-to-oranges and is intentionally not the headline of the paper.",
    )

    out = OUT_DIR / "table_s1_rf_baseline.docx"
    doc.save(str(out))
    print(f"  written: {out}")


# ──────────────────────────────────────────────────────────────────────
# Table S2 — Per-stratum coverage
# ──────────────────────────────────────────────────────────────────────


def render_table_s2():
    data = json.loads((SUP_DATA / "table_s2_per_stratum_coverage.json").read_text())
    doc = Document()
    _apply_page_layout(doc)

    _add_heading(doc, "Supplementary Table S2", level=1)
    _add_caption(
        doc,
        "Table S2.",
        "Per-target, per-stratum empirical conformal coverage on the held-out test split "
        f"(nominal = {data['_meta']['nominal_coverage']:.2f}, α = {data['_meta']['alpha']:.2f}). "
        "For each (target, stratum) cell the table reports the empirical coverage rate (point), "
        "the Clopper–Pearson exact 95 % binomial CI (CP), the 95 % paired-bootstrap percentile CI "
        f"on the coverage rate ({data['_meta']['n_boot']} resamples), and the per-stratum sample "
        "size n. The CQR row at the bottom is the heteroscedastic Conformalized Quantile Regression "
        "layer that supersedes the homoscedastic Mondrian baseline on `time_to_gloc_s` under OSF "
        "amendment 2026-05-06 hypothesis H5. Per-stratum bootstrap CIs for the CQR row are not "
        "computed here because the upstream CQR artifact only commits the per-stratum point rates; "
        "the overall CQR coverage carries the Clopper–Pearson CI on n = 36 quoted in §3.3.",
    )

    headers = [
        "Target / stage",
        "Stratum",
        "n",
        "k",
        "Coverage",
        "CP 95 % CI",
        "Bootstrap 95 % CI",
    ]
    aligns = [
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
    ]
    rows = []
    for r in data["rows"]:
        # Overall row
        o = r["overall"]
        rows.append([
            r["target"],
            "Overall",
            str(o["n"]),
            str(o["k"]),
            _fmt_pct(o["coverage"]),
            _fmt_cp(o["cp_lo"], o["cp_hi"]),
            _fmt_cp(o.get("bs_lo"), o.get("bs_hi")),
        ])
        # Per stratum
        for s in ("championship", "conceptual", "extreme_post_stall", "military_acm"):
            d = r["strata"][s]
            rows.append([
                "",
                "  " + s.replace("_", " "),
                str(d["n"]),
                str(d["k"]),
                _fmt_pct(d["coverage"]),
                _fmt_cp(d["cp_lo"], d["cp_hi"]),
                _fmt_cp(d.get("bs_lo"), d.get("bs_hi")),
            ])

    col_widths = [5.0, 3.2, 1.0, 1.0, 1.8, 2.6, 2.8]
    _build_table(doc, headers, rows, col_widths_cm=col_widths, align_per_col=aligns)

    _add_footnote(
        doc,
        "Footnote — n is the per-stratum test sample size after stratified split (seed 42); k is "
        "the count of in-interval rows. Strata with n = 0 (e.g., conceptual and extreme_post_stall "
        "for `time_to_gloc_s` regressor — no event-positive rows in those strata of the OSF test "
        "slice) are reported as 0/0 with NaN coverage; this is a function of the dataset structure, "
        "not of the conformal layer. Strata with n < 20 carry binomial CIs wider than ±10 pp and "
        "are not reliable for a per-stratum claim; the operationally meaningful stratum that drives "
        "`time_to_gloc_s` overall coverage is military_acm (n = 35). For the CQR row, per-stratum "
        "n values are reproduced from the committed artifact "
        "`data/results/cqr/cqr_vs_mondrian_time_to_gloc.json`; per-stratum bootstrap CIs are not "
        "available because that artifact commits only the per-stratum point rates. Source code: "
        "`scripts/build_supplementary_tables.py`. Calibrator: "
        "`cgem_ext.surrogate.conformal.MondrianSplitConformal` (homoscedastic) and "
        "`cgem_ext.surrogate.cqr.TwoStageXGBQuantileSurrogate` (heteroscedastic CQR).",
    )

    out = OUT_DIR / "table_s2_per_stratum_coverage.docx"
    doc.save(str(out))
    print(f"  written: {out}")


# ──────────────────────────────────────────────────────────────────────
# Table S3 — Second-order Sobol interactions
# ──────────────────────────────────────────────────────────────────────


def render_table_s3():
    """The manuscript supplementary list item 10 — second-order Sobol with
    95 % bootstrap CIs. The committed CSV already carries S2 + S2_conf
    (SALib's bootstrap half-width); we render the top |S2| pairs per
    target as a DOCX table."""
    if not SOBOL2.exists():
        print(f"  skip Table S3 — {SOBOL2} not present")
        return
    by_target: dict[str, list[dict]] = {}
    with SOBOL2.open() as fh:
        for row in csv.DictReader(fh):
            by_target.setdefault(row["target"], []).append(row)

    doc = Document()
    _apply_page_layout(doc)
    _add_heading(doc, "Supplementary Table S3", level=1)
    _add_caption(
        doc,
        "Table S3.",
        "Second-order Sobol interaction indices (S₂) for the top |S₂| pairs per target, with the "
        "SALib bootstrap half-width (95 % CI: S₂ ± S2_conf). All five surrogate targets are "
        "included. Negative S₂ values within ±2 × half-width are consistent with sampling noise at "
        "the N = 1,024 Saltelli base sample. Rows are sorted by |S₂| descending within each target. "
        "Source: `data/results/sensitivity/sobol_second_order.csv`.",
    )

    headers = ["Target", "Feature i", "Feature j", "S₂", "S₂ 95 % CI half-width"]
    aligns = [
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
    ]
    rows: list[list[str]] = []
    for target, items in by_target.items():
        items_sorted = sorted(items, key=lambda r: -abs(float(r["S2"])))
        top = items_sorted[:6]
        for k, row in enumerate(top):
            rows.append([
                target if k == 0 else "",
                row["feature_i"],
                row["feature_j"],
                f"{float(row['S2']):+.4f}",
                f"{float(row['S2_conf']):.4f}",
            ])

    col_widths = [3.6, 3.6, 3.6, 2.4, 2.4]
    _build_table(doc, headers, rows, col_widths_cm=col_widths, align_per_col=aligns)

    _add_footnote(
        doc,
        "Footnote — Sobol indices are computed via SALib's Saltelli sampler (N = 1,024 base, "
        "20,480 total surrogate evaluations) and standard Sobol estimator. Bootstrap half-widths "
        "are SALib's default. The strongest interaction across all time targets is "
        "`g_peak_abs × profile_duration_s`, consistent with the discussion in §3.6.",
    )

    out = OUT_DIR / "table_s3_second_order_sobol.docx"
    doc.save(str(out))
    print(f"  written: {out}")


# ──────────────────────────────────────────────────────────────────────
# Supplementary Information cover doc
# ──────────────────────────────────────────────────────────────────────


def render_cover():
    doc = Document()
    _apply_page_layout(doc)
    _add_heading(doc, "Supplementary Information", level=1)
    _add_para(
        doc,
        "Manuscript title: Calibrated surrogate emulator, conformal prediction intervals, "
        "and global sensitivity analysis for the FAA CGEM G-effects model — a portable "
        "uncertainty-aware extension layer for a validated regulatory ODE physiological "
        "model.",
    )
    _add_para(doc, "Author: Diego Malpica, MD. Sole author.", italic=True)
    _add_para(
        doc,
        "Journal: International Journal for Numerical Methods in Biomedical Engineering "
        "(IJNMBE), Wiley.",
        italic=True,
    )

    _add_heading(doc, "Portal-upload mapping (Wiley CNM portal)", level=2)
    _add_para(
        doc,
        "Per the IJNMBE Author Guidelines, primary data and code are uploaded as "
        "Data Files (a distinct portal designation from Supporting Information). "
        "On acceptance Wiley deposits Data Files to figshare under CC-Zero and "
        "assigns a single DOI permanently linked to the article HTML. The "
        "supplementary package is therefore split into two subfolders:",
    )
    _add_para(
        doc,
        "    supporting_info/   ← upload as Wiley 'Supporting Information' (DOCX tables, "
        "model cards, datasheet, TRIPOD-AI checklist, OSF pre-registration + amendment).",
    )
    _add_para(
        doc,
        "    data_files/        ← upload as Wiley 'Data Files' (code, raw JSON / CSV / "
        "Parquet artefacts behind every table; SHAP and Morris plots; H6 evaluation data).",
    )

    _add_heading(doc, "Index of Supplementary Materials", level=2)
    _add_para(
        doc,
        "Reference column denotes the manuscript-citation key (Table Sx / Document Sx / "
        "Data Sx / Code Sx). File column gives the on-disk filename in the relevant "
        "subfolder above.",
    )

    items = [
        ("Table S1",    "RandomForest baseline vs XGBoost surrogate on the OSF-pre-registered held-out test split. Cited from manuscript §3.2.",                                       "table_s1_rf_baseline.docx"),
        ("Table S2",    "Per-target, per-stratum empirical conformal coverage with Clopper–Pearson exact 95 % binomial CIs and bootstrap-resampled coverage CIs. Cited from manuscript §3.3.", "table_s2_per_stratum_coverage.docx"),
        ("Table S3",    "Second-order Sobol interaction indices (S₂) with SALib bootstrap half-widths. Manuscript-§supp item 10.",                                                     "table_s3_second_order_sobol.docx"),
        ("Document S1", "TRIPOD-AI compliance checklist (manuscript-§supp item 1).",                                                                                                  "tripod_ai_checklist.docx"),
        ("Document S2", "Dataset datasheet, Gebru-style (manuscript-§supp item 2).",                                                                                                   "datasheet.docx"),
        ("Document S3", "Emulator model card (manuscript-§supp item 3).",                                                                                                              "emulator_card.docx"),
        ("Document S4", "OOD detector model card (manuscript-§supp item 4).",                                                                                                          "ood_card.docx"),
        ("Document S5", "OSF pre-registration (master text, locking split indices and success thresholds).",                                                                            "osf_preregistration.docx"),
        ("Document S6", "OSF amendment 2026-05-06 (H5 CQR anchor; H6 archival validation cohort).",                                                                                     "osf_amendment_2026-05-06.docx"),
        ("Data S1",     "Frozen OSF-pre-registered hyperparameter search spaces (JSON, manuscript-§supp item 5).",                                                                       "osf_search_spaces.json"),
        ("Data S2",     "Frozen OSF-pre-registered split indices (Parquet, manuscript-§supp item 6).",                                                                                  "osf_split_indices.parquet"),
        ("Data S3",     "SHAP TreeExplainer mean(|SHAP|) per (target, feature) on the held-out test split (manuscript-§supp item 8).",                                                  "shap_importance.json"),
        ("Data S4",     "Morris Elementary Effects μ* and σ per (target, feature) (manuscript-§supp item 9, CSV).",                                                                      "morris.csv"),
        ("Data S5",     "First and total-order Sobol indices (CSV) — also rendered as manuscript Figure 5.",                                                                            "sobol_first_total.csv"),
        ("Data S6",     "Second-order Sobol interactions raw CSV (also rendered as Table S3).",                                                                                          "sobol_second_order.csv"),
        ("Code S1",     "Reproducibility runner for Tables S1–S2 and the JSON artifacts behind them (Python).",                                                                          "build_supplementary_tables.py"),
        ("Code S2",     "DOCX renderer for Tables S1–S3 and this Supplementary Information cover (Python).",                                                                             "render_supplementary_docx.py"),
        ("Code S3",     "OSF artifact materializer — emits osf_search_spaces.json and osf_split_indices.parquet (Python).",                                                              "build_osf_artifacts.py"),
        ("Code S4",     "SHAP TreeExplainer driver — emits shap_importance.json (Python).",                                                                                              "build_shap_supplementary.py"),
        ("Code S5",     "H6 archival validation evaluation script (Python).",                                                                                                            "run_h6_evaluation.py"),
    ]

    headers = ["Reference", "Description", "File"]
    aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT]
    rows = [[ref, desc, fname] for ref, desc, fname in items]
    col_widths = [2.4, 9.6, 4.0]
    _build_table(doc, headers, rows, col_widths_cm=col_widths, align_per_col=aligns)

    _add_heading(doc, "Reproducibility notes", level=2)
    _add_para(
        doc,
        "All numeric supplementary tables are reproducible from the frozen dataset "
        "`data/datasets/cgem_synthetic_v1.parquet` (SHA-256 and master seed recorded in the "
        "sidecar `cgem_synthetic_v1.meta.json`) by running:",
    )
    _add_para(
        doc,
        "    python scripts/build_supplementary_tables.py",
    )
    _add_para(
        doc,
        "    python scripts/render_supplementary_docx.py",
    )
    _add_para(
        doc,
        "Source code is open under the MIT licence at "
        "https://github.com/strikerdlm/CAMI-Gz-Effects-Model-CGEM- and the dataset is "
        "archived on Zenodo (DOI: TBD at submission) per the Joint Declaration of Data "
        "Citation Principles. The data and code are uploaded to the Wiley CNM portal as "
        "Data Files; the items in the table above are uploaded as Supporting Information.",
    )

    out = OUT_DIR / "supplementary_information.docx"
    doc.save(str(out))
    print(f"  written: {out}")


# ──────────────────────────────────────────────────────────────────────
# Render the markdown -> DOCX items (TRIPOD-AI, datasheet, model cards)
# via pandoc using the IJNMBE reference DOCX.
# ──────────────────────────────────────────────────────────────────────


def render_md_items():
    import shutil
    import subprocess

    ref = REPO / "manuscripts" / "ijnmbe" / "rendered" / "_reference.docx"
    sources = [
        (REPO / "docs/publication/tripod_ai_checklist.md", "tripod_ai_checklist.docx"),
        (REPO / "docs/data/datasheet.md",                   "datasheet.docx"),
        (REPO / "docs/models/emulator_card.md",             "emulator_card.docx"),
        (REPO / "docs/models/ood_card.md",                  "ood_card.docx"),
        (REPO / "docs/publication/osf_preregistration.md",  "osf_preregistration.docx"),
        (REPO / "docs/publication/osf_amendment_2026-05-06.md", "osf_amendment_2026-05-06.docx"),
    ]
    if shutil.which("pandoc") is None:
        print("  WARNING: pandoc not on PATH — skipping markdown renders.")
        return
    for src, dst_name in sources:
        if not src.is_file():
            print(f"  MISSING: {src} (skipped)")
            continue
        dst = OUT_DIR / dst_name
        cmd = ["pandoc", str(src), "-o", str(dst), "--standalone", "--wrap=none"]
        if ref.is_file():
            cmd += [f"--reference-doc={ref}"]
        subprocess.run(cmd, check=True, capture_output=True)
        print(f"  pandoc → supporting_info/{dst_name}")


def stage_data_files():
    """Copy raw data + code into the data_files subdirectory."""
    import shutil

    sources = [
        # Frozen OSF artifacts
        (REPO / "docs/publication/osf_search_spaces.json",       "osf_search_spaces.json"),
        (REPO / "docs/publication/osf_split_indices.parquet",    "osf_split_indices.parquet"),
        # SHAP / sensitivity raw data
        (REPO / "data/results/supplementary/shap_importance.json", "shap_importance.json"),
        (REPO / "data/results/sensitivity/morris.csv",             "morris.csv"),
        (REPO / "data/results/sensitivity/sobol_first_total.csv",  "sobol_first_total.csv"),
        (REPO / "data/results/sensitivity/sobol_second_order.csv", "sobol_second_order.csv"),
        # H6 artifacts
        (REPO / "data/results/h6/discrepancy_phase_a.json",        "h6_discrepancy_phase_a.json"),
        (REPO / "data/results/h6/multifidelity_benchmark.json",    "h6_multifidelity_benchmark.json"),
        # CQR vs Mondrian artifact (also referenced from S2)
        (REPO / "data/results/cqr/cqr_vs_mondrian_time_to_gloc.json", "cqr_vs_mondrian_time_to_gloc.json"),
        # Underlying Table S1 / S2 JSON
        (REPO / "data/results/supplementary/table_s1_rf_baseline.json",        "table_s1_rf_baseline.json"),
        (REPO / "data/results/supplementary/table_s2_per_stratum_coverage.json","table_s2_per_stratum_coverage.json"),
        # Code
        (REPO / "scripts/build_supplementary_tables.py",     "build_supplementary_tables.py"),
        (REPO / "scripts/render_supplementary_docx.py",      "render_supplementary_docx.py"),
        (REPO / "scripts/build_osf_artifacts.py",            "build_osf_artifacts.py"),
        (REPO / "scripts/build_shap_supplementary.py",       "build_shap_supplementary.py"),
        (REPO / "scripts/build_shap_morris_plots.py",        "build_shap_morris_plots.py"),
        (REPO / "scripts/run_h6_evaluation.py",              "run_h6_evaluation.py"),
        (REPO / "scripts/generate_figure_data.py",           "generate_figure_data.py"),
    ]
    for src, dst_name in sources:
        if not src.is_file():
            print(f"  MISSING (skipped): {src}")
            continue
        shutil.copy2(src, DATA_DIR / dst_name)
        print(f"  copied → data_files/{dst_name}")

    # Plots directory (mirror)
    plots_src = REPO / "data/results/supplementary/plots"
    if plots_src.is_dir():
        plots_dst = DATA_DIR / "plots"
        plots_dst.mkdir(exist_ok=True)
        for p in plots_src.iterdir():
            shutil.copy2(p, plots_dst / p.name)
        print(f"  copied {len(list(plots_src.iterdir()))} plot files → data_files/plots/")


# ──────────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────────


def main() -> int:
    print(f"Supporting Information dir: {OUT_DIR}")
    print(f"Data Files dir:             {DATA_DIR}")
    print("\n=== Rendering Table S1 ===")
    render_table_s1()
    print("\n=== Rendering Table S2 ===")
    render_table_s2()
    print("\n=== Rendering Table S3 ===")
    render_table_s3()
    print("\n=== Rendering markdown items (TRIPOD-AI, datasheet, model cards, OSF) ===")
    render_md_items()
    print("\n=== Rendering Supplementary Information cover ===")
    render_cover()
    print("\n=== Staging Data Files (code + raw artefacts) ===")
    stage_data_files()
    print(f"\nDone.")
    print(f"  Upload contents of {OUT_DIR} as Wiley 'Supporting Information'.")
    print(f"  Upload contents of {DATA_DIR} as Wiley 'Data Files'.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
