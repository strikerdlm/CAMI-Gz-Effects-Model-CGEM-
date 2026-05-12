"""Generate a Wiley Free-Format-compliant pandoc reference DOCX for IJNMBE.

Formatting targets (standard peer-review manuscript):
  - Font:         Times New Roman 12 pt (body)
  - Line spacing: Double (exact 24 pt)
  - Line numbers: Continuous, restart each page, left margin
  - Margins:      2.54 cm (1 in) all sides
  - Page size:    A4 (21 × 29.7 cm)
  - Page numbers: Top-right footer (header), starting at 1
  - Headings:     Bold, same font, sized H1=14/H2=13/H3=12
  - Ragged right: left-aligned body text (not justified)

Usage:
    python scripts/build_ijnmbe_reference_docx.py
Outputs:
    manuscripts/ijnmbe/rendered/_reference.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
OUT = REPO / "manuscripts" / "ijnmbe" / "rendered" / "_reference.docx"


def _set_font(style, name: str, size_pt: float, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    # Force East-Asian and complex-script fonts too
    rPr = style.element.get_or_add_rPr()
    for tag in ("w:rFonts",):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:ascii"), name)
        el.set(qn("w:hAnsi"), name)
        el.set(qn("w:cs"), name)
        el.set(qn("w:eastAsia"), name)


def _set_double_spacing(style) -> None:
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_line_numbering(section) -> None:
    """Add continuous line numbering to a section via raw OOXML."""
    sectPr = section._sectPr
    lnNumType = OxmlElement("w:lnNumType")
    lnNumType.set(qn("w:countBy"), "1")
    lnNumType.set(qn("w:restart"), "newPage")
    lnNumType.set(qn("w:start"), "1")
    lnNumType.set(qn("w:distance"), "720")  # 0.5 in from text
    sectPr.append(lnNumType)


def _set_page(section) -> None:
    """A4, 2.54 cm margins."""
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)


def _add_page_number_header(section) -> None:
    """Add right-aligned page number in the header."""
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def main() -> None:
    doc = Document()

    # --- Section / page layout ---
    section = doc.sections[0]
    _set_page(section)
    _add_line_numbering(section)
    _add_page_number_header(section)

    # --- Normal (body) style ---
    normal = doc.styles["Normal"]
    _set_font(normal, "Times New Roman", 12)
    _set_double_spacing(normal)

    # --- Heading styles ---
    for level, size, bold in [(1, 14, True), (2, 13, True), (3, 12, True)]:
        try:
            h = doc.styles[f"Heading {level}"]
        except KeyError:
            h = doc.styles.add_style(f"Heading {level}", 1)
        _set_font(h, "Times New Roman", size, bold=bold)
        pf = h.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- Caption / Table caption style ---
    for sname in ("Caption", "Table Caption", "Figure Caption"):
        try:
            cap = doc.styles[sname]
        except KeyError:
            continue
        _set_font(cap, "Times New Roman", 11)
        cap.paragraph_format.space_before = Pt(6)
        cap.paragraph_format.space_after = Pt(6)

    # --- Block text / body text styles ---
    for sname in ("Body Text", "Body Text 2", "Body Text 3"):
        try:
            bt = doc.styles[sname]
            _set_font(bt, "Times New Roman", 12)
            _set_double_spacing(bt)
        except KeyError:
            continue

    # --- First paragraph style (no indent) ---
    try:
        fp = doc.styles["First Paragraph"]
        _set_font(fp, "Times New Roman", 12)
        _set_double_spacing(fp)
    except KeyError:
        pass

    # --- Table styles ---
    try:
        ts = doc.styles["Table Grid"]
        ts.font.name = "Times New Roman"
        ts.font.size = Pt(11)
    except KeyError:
        pass

    # Add a placeholder paragraph so the file is non-empty
    p = doc.add_paragraph("Reference document — Wiley IJNMBE submission (Free Format)")
    p.style = doc.styles["Normal"]

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Written: {OUT}")


if __name__ == "__main__":
    main()
